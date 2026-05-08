#!/usr/bin/env python3
"""Smoke test for the inline subscript normalization script."""

from __future__ import annotations

import json
import subprocess
import sys
import tempfile
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts" / "unify_inline_subscripts.py"


def main() -> int:
    if not SCRIPT.exists():
        raise AssertionError(f"missing script: {SCRIPT}")

    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        input_docx = tmpdir / "input.docx"
        output_docx = tmpdir / "output.docx"
        config_json = tmpdir / "config.json"

        doc = Document()
        doc.add_paragraph("模型输入变量中的f_WC、x_Co、x_Ni和T_norm需要统一。")
        doc.add_paragraph("这段不应该被修改。")
        doc.save(input_docx)

        config = {
            "paragraph_indices": [1],
            "tokens": [
                {"token": "f_WC", "base": "f", "sub": "WC"},
                {"token": "x_Co", "base": "x", "sub": "Co"},
                {"token": "x_Ni", "base": "x", "sub": "Ni"},
                {"token": "T_norm", "base": "T", "sub": "norm"},
            ],
        }
        config_json.write_text(json.dumps(config, ensure_ascii=False), encoding="utf-8")

        result = subprocess.run(
            [sys.executable, str(SCRIPT), str(input_docx), str(output_docx), str(config_json)],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=False,
        )
        if result.returncode != 0:
            raise AssertionError(result.stderr or result.stdout)

        with ZipFile(output_docx) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")

        if "f_WC" in document_xml or "x_Co" in document_xml or "T_norm" in document_xml:
            raise AssertionError("raw underscore tokens should not remain in output XML")
        if document_xml.count('w:vertAlign w:val="subscript"') < 4:
            raise AssertionError("expected at least four subscript runs")

    print("PASS test_unify_inline_subscripts")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
