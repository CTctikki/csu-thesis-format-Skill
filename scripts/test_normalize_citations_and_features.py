#!/usr/bin/env python3
"""Smoke test for citation superscripts and inline feature notation."""

from __future__ import annotations

import subprocess
import sys
import tempfile
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts" / "normalize_citations_and_features.py"


def xml_has_vert(document_xml: str, text: str, value: str) -> bool:
    marker = f"<w:t>{text}</w:t>"
    idx = document_xml.find(marker)
    if idx < 0:
        return False
    run_start = document_xml.rfind("<w:r", 0, idx)
    run_end = document_xml.find("</w:r>", idx)
    run_xml = document_xml[run_start:run_end]
    return f'w:vertAlign w:val="{value}"' in run_xml


def main() -> int:
    if not SCRIPT.exists():
        raise AssertionError(f"missing script: {SCRIPT}")

    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        input_docx = tmpdir / "input.docx"
        output_docx = tmpdir / "output.docx"

        doc = Document()
        doc.add_paragraph("已有研究表明热导率受dWC、fWC和xNi影响[1-3]。")
        doc.add_paragraph("派生特征包括aWC、Asp、Tnorm、fWC2/3、R2和cm-3。")
        doc.add_paragraph("参考文献")
        doc.add_paragraph("[1] WEN S, TAN J. Example reference[J].")
        doc.save(input_docx)

        result = subprocess.run(
            [sys.executable, str(SCRIPT), str(input_docx), str(output_docx)],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=False,
        )
        if result.returncode != 0:
            raise AssertionError(result.stderr or result.stdout)

        with ZipFile(output_docx) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")

        if not xml_has_vert(document_xml, "[1-3]", "superscript"):
            raise AssertionError("body citation should be superscript")
        if xml_has_vert(document_xml, "[1] WEN S, TAN J. Example reference[J].", "superscript"):
            raise AssertionError("reference-list entry should not be superscripted")
        for text in ["WC", "Ni", "sp", "norm"]:
            if not xml_has_vert(document_xml, text, "subscript"):
                raise AssertionError(f"expected subscript run for {text}")
        for text in ["2/3", "2", "-3"]:
            if not xml_has_vert(document_xml, text, "superscript"):
                raise AssertionError(f"expected superscript run for {text}")

    print("PASS test_normalize_citations_and_features")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
