#!/usr/bin/env python3
"""Extract a compact formatting inventory from a Word DOCX file."""

from __future__ import annotations

import argparse
from collections import Counter
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}


def qn(tag: str) -> str:
    return f"{{{W}}}{tag}"


def attr(el, name: str = "val"):
    return el.get(qn(name)) if el is not None else None


def twips_to_cm(value) -> float | None:
    return round(int(value) / 1440 * 2.54, 3) if value is not None else None


def twips_to_pt(value) -> float | None:
    return round(int(value) / 20, 2) if value is not None else None


def half_points_to_pt(value) -> float | None:
    return round(int(value) / 2, 2) if value is not None else None


def spacing_desc(spacing) -> dict:
    if spacing is None:
        return {}
    data = {}
    for key in ["before", "after", "line", "lineRule"]:
        value = spacing.get(qn(key))
        if value is not None:
            data[key] = value
    if "before" in data:
        data["before_pt"] = twips_to_pt(data["before"])
    if "after" in data:
        data["after_pt"] = twips_to_pt(data["after"])
    if "line" in data:
        if data.get("lineRule") in {"exact", "atLeast"}:
            data["line_pt"] = twips_to_pt(data["line"])
        else:
            data["line_multiple"] = round(int(data["line"]) / 240, 3)
    return data


def indent_desc(indent) -> dict:
    if indent is None:
        return {}
    data = {}
    for key in ["left", "right", "firstLine", "hanging", "firstLineChars", "hangingChars"]:
        value = indent.get(qn(key))
        if value is not None:
            data[key] = value
    for key in ["left", "right", "firstLine", "hanging"]:
        if key in data:
            data[f"{key}_cm"] = twips_to_cm(data[key])
    return data


def paragraph_props(p_pr) -> dict:
    if p_pr is None:
        return {}
    data = {}
    style = attr(p_pr.find("w:pStyle", NS))
    if style:
        data["style_id"] = style
    align = attr(p_pr.find("w:jc", NS))
    if align:
        data["align"] = align
    spacing = spacing_desc(p_pr.find("w:spacing", NS))
    if spacing:
        data["spacing"] = spacing
    indent = indent_desc(p_pr.find("w:ind", NS))
    if indent:
        data["indent"] = indent
    outline = attr(p_pr.find("w:outlineLvl", NS))
    if outline is not None:
        data["outline_level"] = outline
    return data


def run_props(r_pr) -> dict:
    if r_pr is None:
        return {}
    data = {}
    fonts = r_pr.find("w:rFonts", NS)
    if fonts is not None:
        font_data = {}
        for key in ["ascii", "hAnsi", "eastAsia", "cs", "hint"]:
            value = fonts.get(qn(key))
            if value:
                font_data[key] = value
        if font_data:
            data["fonts"] = font_data
    size = attr(r_pr.find("w:sz", NS))
    if size:
        data["size_pt"] = half_points_to_pt(size)
    size_cs = attr(r_pr.find("w:szCs", NS))
    if size_cs:
        data["size_cs_pt"] = half_points_to_pt(size_cs)
    color = attr(r_pr.find("w:color", NS))
    if color:
        data["color"] = color
    for tag in ["b", "i", "u", "smallCaps", "caps", "strike"]:
        el = r_pr.find(f"w:{tag}", NS)
        if el is not None:
            data[tag] = el.get(qn("val"), "1")
    return data


def load_xml(path: Path, member: str):
    with ZipFile(path) as archive:
        return etree.fromstring(archive.read(member))


def style_definitions(path: Path) -> dict:
    styles = load_xml(path, "word/styles.xml")
    result = {}
    for style in styles.findall("w:style", NS):
        name = attr(style.find("w:name", NS)) or style.get(qn("styleId"))
        result[name] = {
            "style_id": style.get(qn("styleId")),
            "type": style.get(qn("type")),
            "based_on": attr(style.find("w:basedOn", NS)),
            "next": attr(style.find("w:next", NS)),
            "paragraph": paragraph_props(style.find("w:pPr", NS)),
            "run": run_props(style.find("w:rPr", NS)),
        }
    return result


def xml_paragraph_samples(path: Path, limit: int) -> list[dict]:
    document = load_xml(path, "word/document.xml")
    samples = []
    for paragraph in document.findall(".//w:p", NS):
        text = "".join(t.text or "" for t in paragraph.findall(".//w:t", NS)).strip()
        if not text:
            continue
        p_pr = paragraph.find("w:pPr", NS)
        r_pr = None
        for run in paragraph.findall("w:r", NS):
            run_text = "".join(t.text or "" for t in run.findall(".//w:t", NS)).strip()
            if run_text:
                r_pr = run.find("w:rPr", NS)
                break
        samples.append(
            {
                "text": text[:120],
                "paragraph": paragraph_props(p_pr),
                "run": run_props(r_pr),
            }
        )
        if len(samples) >= limit:
            break
    return samples


def print_dict(data: dict, indent: str = "  ") -> None:
    for key, value in data.items():
        print(f"{indent}- {key}: {value}")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--samples", type=int, default=40, help="number of XML paragraph samples")
    args = parser.parse_args()

    docx_path = args.docx
    if not docx_path.exists():
        raise SystemExit(f"File not found: {docx_path}")

    document = Document(str(docx_path))
    print(f"# DOCX Format Report: {docx_path.name}\n")

    print("## Sections")
    for index, section in enumerate(document.sections, 1):
        print(f"- Section {index}:")
        print(f"  - page: {round(section.page_width.cm, 3)} cm x {round(section.page_height.cm, 3)} cm")
        print(
            "  - margins: "
            f"top {round(section.top_margin.cm, 3)} cm, "
            f"bottom {round(section.bottom_margin.cm, 3)} cm, "
            f"left {round(section.left_margin.cm, 3)} cm, "
            f"right {round(section.right_margin.cm, 3)} cm"
        )
        print(
            "  - header/footer: "
            f"header {round(section.header_distance.cm, 3)} cm, "
            f"footer {round(section.footer_distance.cm, 3)} cm, "
            f"different_first_page={section.different_first_page_header_footer}"
        )

    body_counts = Counter(p.style.name for p in document.paragraphs if p.text.strip())
    print("\n## Body Paragraph Styles")
    for name, count in body_counts.most_common():
        print(f"- {name}: {count}")

    print(f"\n## Tables\n- Count: {len(document.tables)}")
    for index, table in enumerate(document.tables[:10], 1):
        rows = len(table.rows)
        cols = len(table.columns)
        preview = " | ".join(" ".join(cell.text.split()) for cell in table.rows[0].cells) if rows else ""
        print(f"- Table {index}: {rows} rows x {cols} columns; first row: {preview[:120]}")

    styles = style_definitions(docx_path)
    print("\n## Used Style Definitions")
    for name in body_counts:
        print(f"- {name}")
        print_dict(styles.get(name, {"note": "style not found in styles.xml"}), "  ")

    print("\n## XML Paragraph Samples")
    for index, sample in enumerate(xml_paragraph_samples(docx_path, args.samples), 1):
        print(f"- Sample {index}: {sample['text']}")
        if sample["paragraph"]:
            print(f"  - paragraph: {sample['paragraph']}")
        if sample["run"]:
            print(f"  - run: {sample['run']}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
