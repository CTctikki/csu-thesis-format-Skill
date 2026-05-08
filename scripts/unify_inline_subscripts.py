#!/usr/bin/env python3
"""Normalize inline variable subscripts in selected DOCX paragraphs."""

from __future__ import annotations

import argparse
import copy
import json
from dataclasses import dataclass
from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


@dataclass(frozen=True)
class TokenSpec:
    token: str
    base: str
    sub: str


def load_config(config_path: Path) -> tuple[list[int], list[TokenSpec]]:
    config = json.loads(config_path.read_text(encoding="utf-8"))
    indices = set(config.get("paragraph_indices", []))
    for item in config.get("paragraph_ranges", []):
        start = int(item["start"])
        end = int(item["end"])
        indices.update(range(start, end + 1))

    specs = [
        TokenSpec(str(item["token"]), str(item["base"]), str(item["sub"]))
        for item in config.get("tokens", [])
    ]
    specs.sort(key=lambda item: len(item.token), reverse=True)
    return sorted(indices), specs


def first_run_rpr(paragraph):
    for run in paragraph.runs:
        if run.text or run._r.rPr is not None:
            return run._r.rPr
    return None


def clone_rpr(base_rpr, subscript: bool):
    if base_rpr is None:
        rpr = OxmlElement("w:rPr")
    else:
        rpr = copy.deepcopy(base_rpr)
    vert = rpr.find(qn("w:vertAlign"))
    if vert is not None:
        rpr.remove(vert)
    if subscript:
        vert = OxmlElement("w:vertAlign")
        vert.set(qn("w:val"), "subscript")
        rpr.append(vert)
    return rpr


def remove_non_ppr_children(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def build_parts(text: str, specs: list[TokenSpec]) -> list[tuple[str, bool]]:
    parts: list[tuple[str, bool]] = []
    i = 0
    while i < len(text):
        matched = False
        for spec in specs:
            if text.startswith(spec.token, i):
                parts.append((spec.base, False))
                parts.append((spec.sub, True))
                i += len(spec.token)
                matched = True
                break
        if matched:
            continue

        start = i
        while i < len(text) and not any(text.startswith(spec.token, i) for spec in specs):
            i += 1
        parts.append((text[start:i], False))
    return parts


def append_run(paragraph, text: str, rpr) -> None:
    if not text:
        return
    run = OxmlElement("w:r")
    run.append(rpr)
    t = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace() or "  " in text:
        t.set(XML_SPACE, "preserve")
    t.text = text
    run.append(t)
    paragraph._p.append(run)


def normalize_paragraph(paragraph, specs: list[TokenSpec]) -> bool:
    text = paragraph.text
    if not any(spec.token in text for spec in specs):
        return False

    base_rpr = first_run_rpr(paragraph)
    normal_rpr = clone_rpr(base_rpr, subscript=False)
    sub_rpr = clone_rpr(base_rpr, subscript=True)
    parts = build_parts(text, specs)

    remove_non_ppr_children(paragraph)
    for chunk, is_sub in parts:
        append_run(paragraph, chunk, copy.deepcopy(sub_rpr if is_sub else normal_rpr))
    return True


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("config_json", type=Path)
    args = parser.parse_args()

    if not args.input_docx.exists():
        raise SystemExit(f"missing input DOCX: {args.input_docx}")
    if not args.config_json.exists():
        raise SystemExit(f"missing config JSON: {args.config_json}")

    if args.input_docx.resolve() != args.output_docx.resolve():
        copyfile(args.input_docx, args.output_docx)

    paragraph_indices, specs = load_config(args.config_json)
    doc = Document(str(args.output_docx))

    changed = 0
    for index in paragraph_indices:
        if 1 <= index <= len(doc.paragraphs):
            changed += int(normalize_paragraph(doc.paragraphs[index - 1], specs))

    doc.save(str(args.output_docx))
    print(f"[OK] saved {args.output_docx}")
    print(f"[OK] normalized paragraphs: {changed}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
