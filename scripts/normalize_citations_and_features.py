#!/usr/bin/env python3
"""Normalize thesis citation superscripts and compact feature notation in DOCX."""

from __future__ import annotations

import argparse
import re
from copy import deepcopy
from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
PRESERVABLE_RUN_PAYLOAD = {"lastRenderedPageBreak"}

CITATION_RE = re.compile(r"\[[0-9]+(?:\s*[-,，－–—]\s*[0-9]+)*\]")

FEATURE_TOKENS: dict[str, list[tuple[str, str | None]]] = {
    "kbinderest": [("k", None), ("binder", "subscript"), ("est", "superscript")],
    "k_binderest": [("k", None), ("binder", "subscript"), ("est", "superscript")],
    "k_binder_est": [("k", None), ("binder", "subscript"), ("est", "superscript")],
    "fWC2/3": [("f", None), ("WC", "subscript"), ("2/3", "superscript")],
    "f_WC2/3": [("f", None), ("WC", "subscript"), ("2/3", "superscript")],
    "k300K": [("k", None), ("300K", "subscript")],
    "k1100K": [("k", None), ("1100K", "subscript")],
    "dWC": [("d", None), ("WC", "subscript")],
    "d_WC": [("d", None), ("WC", "subscript")],
    "fWC": [("f", None), ("WC", "subscript")],
    "f_WC": [("f", None), ("WC", "subscript")],
    "xNi": [("x", None), ("Ni", "subscript")],
    "x_Ni": [("x", None), ("Ni", "subscript")],
    "xCo": [("x", None), ("Co", "subscript")],
    "x_Co": [("x", None), ("Co", "subscript")],
    "xAg": [("x", None), ("Ag", "subscript")],
    "x_Ag": [("x", None), ("Ag", "subscript")],
    "aWC": [("a", None), ("WC", "subscript")],
    "a_WC": [("a", None), ("WC", "subscript")],
    "Asp": [("A", None), ("sp", "subscript")],
    "A_sp": [("A", None), ("sp", "subscript")],
    "Tnorm": [("T", None), ("norm", "subscript")],
    "T_norm": [("T", None), ("norm", "subscript")],
    "F0": [("F", None), ("0", "subscript")],
    "F1": [("F", None), ("1", "subscript")],
    "R2": [("R", None), ("2", "superscript")],
    "mm2/s": [("mm", None), ("2", "superscript"), ("/s", None)],
}

FEATURE_WORD_RE = re.compile(
    "|".join(re.escape(token) for token in sorted(FEATURE_TOKENS, key=len, reverse=True))
)
UNIT_POWER_RE = re.compile(r"(cm|μm|µm|um)(-\d+)")


def set_vert_align_on_rpr(r_pr, value: str | None) -> None:
    for existing in r_pr.findall(qn("w:vertAlign")):
        r_pr.remove(existing)
    if value:
        vert = OxmlElement("w:vertAlign")
        vert.set(qn("w:val"), value)
        r_pr.append(vert)


def make_run_like(source_run, text: str, vert_align: str | None, payload=None):
    run = OxmlElement("w:r")
    source_r_pr = source_run._element.rPr
    r_pr = deepcopy(source_r_pr) if source_r_pr is not None else OxmlElement("w:rPr")
    set_vert_align_on_rpr(r_pr, vert_align)
    run.append(r_pr)

    for child in payload or []:
        run.append(child)

    t = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace() or "  " in text:
        t.set(XML_SPACE, "preserve")
    t.text = text
    run.append(t)
    return run


def unsupported_run_payload(run) -> bool:
    for child in run._element:
        local = child.tag.rsplit("}", 1)[-1]
        if local not in {"rPr", "t", *PRESERVABLE_RUN_PAYLOAD}:
            return True
    return False


def preserved_payload(run):
    payload = []
    for child in run._element:
        local = child.tag.rsplit("}", 1)[-1]
        if local in PRESERVABLE_RUN_PAYLOAD:
            payload.append(deepcopy(child))
    return payload


def feature_match_at(text: str, pos: int):
    word_match = FEATURE_WORD_RE.match(text, pos)
    unit_match = UNIT_POWER_RE.match(text, pos)
    match = None
    if word_match and unit_match:
        match = word_match if len(word_match.group(0)) >= len(unit_match.group(0)) else unit_match
    else:
        match = word_match or unit_match
    if not match:
        return None

    start, end = match.span()
    before = text[start - 1] if start > 0 else ""
    after = text[end] if end < len(text) else ""
    if (before.isascii() and before.isalpha()) or before == "_":
        return None
    if (after.isascii() and after.isalpha()) or after == "_":
        return None
    token = match.group(0)
    if token in FEATURE_TOKENS:
        return end, FEATURE_TOKENS[token]
    unit_base, unit_power = match.groups()
    return end, [(unit_base, None), (unit_power, "superscript")]


def split_segments(text: str, citations: bool, features: bool) -> list[tuple[str, str | None]]:
    parts: list[tuple[str, str | None]] = []
    normal_start = 0
    pos = 0

    while pos < len(text):
        citation = CITATION_RE.match(text, pos) if citations else None
        feature = feature_match_at(text, pos) if features else None
        if citation or feature:
            if normal_start < pos:
                parts.append((text[normal_start:pos], None))
            if citation:
                parts.append((citation.group(0), "superscript"))
                pos = citation.end()
            else:
                end, feature_parts = feature
                parts.extend(feature_parts)
                pos = end
            normal_start = pos
            continue
        pos += 1

    if normal_start < len(text):
        parts.append((text[normal_start:], None))

    merged: list[tuple[str, str | None]] = []
    for chunk, vert in parts:
        if not chunk:
            continue
        if merged and merged[-1][1] == vert:
            merged[-1] = (merged[-1][0] + chunk, vert)
        else:
            merged.append((chunk, vert))
    return merged


def replace_run(run, segments: list[tuple[str, str | None]]) -> bool:
    if unsupported_run_payload(run):
        return False
    if len(segments) == 1 and segments[0] == (run.text, None):
        return False

    parent = run._element.getparent()
    if parent is None:
        return False

    payload = preserved_payload(run)
    insert_at = parent.index(run._element)
    for idx, (text, vert) in enumerate(segments):
        parent.insert(
            insert_at,
            make_run_like(run, text, vert, payload if idx == 0 else None),
        )
        insert_at += 1
    parent.remove(run._element)
    return True


def paragraph_changed(paragraph, citations: bool, features: bool) -> bool:
    changed = False
    for run in list(paragraph.runs):
        if not run.text:
            continue
        segments = split_segments(run.text, citations=citations, features=features)
        changed = replace_run(run, segments) or changed
    return changed


def reference_start_index(doc: Document, heading: str) -> int | None:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == heading:
            return index
    return None


def iter_table_paragraphs(doc: Document):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--reference-heading", default="参考文献")
    parser.add_argument("--no-citations", action="store_true")
    parser.add_argument("--no-features", action="store_true")
    args = parser.parse_args()

    if not args.input_docx.exists():
        raise SystemExit(f"missing input DOCX: {args.input_docx}")
    if args.input_docx.resolve() != args.output_docx.resolve():
        copyfile(args.input_docx, args.output_docx)

    doc = Document(str(args.output_docx))
    ref_start = reference_start_index(doc, args.reference_heading)
    citations = not args.no_citations
    features = not args.no_features

    changed = 0
    for index, paragraph in enumerate(doc.paragraphs):
        in_reference_list = ref_start is not None and index > ref_start
        changed += int(
            paragraph_changed(
                paragraph,
                citations=citations and not in_reference_list,
                features=features and not in_reference_list,
            )
        )

    for paragraph in iter_table_paragraphs(doc):
        changed += int(paragraph_changed(paragraph, citations=citations, features=features))

    doc.save(str(args.output_docx))
    print(f"[OK] saved {args.output_docx}")
    print(f"[OK] normalized paragraphs: {changed}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
