#!/usr/bin/env python3
"""Audit DOCX structures that are hard to reason about from plain text alone."""

from __future__ import annotations

import argparse
import re
from collections import Counter, defaultdict
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL = "http://schemas.openxmlformats.org/package/2006/relationships"
NS = {"w": W, "m": M, "r": R, "rel": PKG_REL}


def qn(tag: str, namespace: str = W) -> str:
    return f"{{{namespace}}}{tag}"


def read_archive_text(archive: ZipFile, member: str) -> str:
    return archive.read(member).decode("utf-8", errors="ignore")


def paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def detect_body_start(document: Document) -> int:
    candidates = []
    for index, paragraph in enumerate(document.paragraphs, 1):
        text = paragraph_text(paragraph)
        if not re.match(r"^第[0-9一二三四五六七八九十百]+章\s+\S", text):
            continue
        if "\t" in text or re.search(r"\.{2,}\s*\d+$|…+\s*\d+$|\s+\d+$", text):
            continue
        candidates.append(index)
    return candidates[0] if candidates else 1


def analyze_sections(document_xml: etree._Element, rel_map: dict[str, str]) -> list[dict]:
    results: list[dict] = []
    for idx, sect in enumerate(document_xml.xpath("//w:sectPr", namespaces=NS), 1):
        headers = []
        footers = []
        for ref in sect.xpath("./w:headerReference", namespaces=NS):
            rid = ref.get(qn("id", R))
            headers.append({"type": ref.get(qn("type")), "rid": rid, "target": rel_map.get(rid, "")})
        for ref in sect.xpath("./w:footerReference", namespaces=NS):
            rid = ref.get(qn("id", R))
            footers.append({"type": ref.get(qn("type")), "rid": rid, "target": rel_map.get(rid, "")})
        pg_num = sect.find("w:pgNumType", NS)
        results.append(
            {
                "index": idx,
                "headers": headers,
                "footers": footers,
                "page_number_format": pg_num.get(qn("fmt")) if pg_num is not None else None,
                "page_number_start": pg_num.get(qn("start")) if pg_num is not None else None,
            }
        )
    return results


def analyze_toc(document: Document, document_xml_text: str, body_start: int) -> dict:
    toc_title = None
    for index, paragraph in enumerate(document.paragraphs, 1):
        if paragraph_text(paragraph) == "目录":
            toc_title = index
            break

    lines: list[dict] = []
    if toc_title is not None:
        for index in range(toc_title + 1, min(body_start, len(document.paragraphs) + 1)):
            paragraph = document.paragraphs[index - 1]
            text = paragraph_text(paragraph)
            if text:
                lines.append({"index": index, "style": paragraph.style.name, "text": text})

    repeated = Counter(text for text in (line["text"] for line in lines) if text.startswith("第"))
    return {
        "has_toc_field": "TOC " in document_xml_text,
        "toc_title_index": toc_title,
        "line_count": len(lines),
        "toc_lines": lines[:80],
        "repeated_headings": {k: v for k, v in repeated.items() if v > 1},
    }


def analyze_headers_and_footers(archive: ZipFile) -> dict:
    header_files = sorted(name for name in archive.namelist() if name.startswith("word/header") and name.endswith(".xml"))
    footer_files = sorted(name for name in archive.namelist() if name.startswith("word/footer") and name.endswith(".xml"))

    header_summaries = []
    for name in header_files:
        xml = read_archive_text(archive, name)
        rel_name = name.replace("word/", "word/_rels/") + ".rels"
        rel_xml = read_archive_text(archive, rel_name) if rel_name in archive.namelist() else ""
        texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml)
        header_summaries.append(
            {
                "name": name,
                "text_preview": " | ".join(t.strip() for t in texts if t.strip())[:200],
                "has_drawing": "<w:drawing" in xml,
                "has_image_rel": "relationships/image" in rel_xml,
                "has_table": "<w:tbl" in xml,
                "has_underline": "<w:u " in xml,
            }
        )

    footer_summaries = []
    for name in footer_files:
        xml = read_archive_text(archive, name)
        texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml)
        footer_summaries.append(
            {
                "name": name,
                "text_preview": " | ".join(t.strip() for t in texts if t.strip())[:200],
                "has_page_field": " PAGE " in xml,
                "has_numpages_field": " NUMPAGES " in xml,
            }
        )
    return {"headers": header_summaries, "footers": footer_summaries}


def looks_like_display_formula(text: str) -> bool:
    if len(text) < 4 or len(text) > 140:
        return False
    if text.startswith(("式中", "其中", "则可通过", "图", "表", "摘要", "目录", "参考文献")):
        return False
    if any(mark in text for mark in ["。", "；", "："]):
        return False
    if re.search(r"[（(]\d+[-－]\d+[）)]\s*$", text):
        return True
    if text.count("=") == 0:
        return False
    return bool(
        re.match(
            r"^[A-Za-z0-9_\-\+\(\)\[\]\|/×÷Σρλα\.,，%\s]+$",
            text,
        )
    )


def analyze_formulas(document: Document, document_xml: etree._Element, body_start: int) -> dict:
    omath_count = len(document_xml.xpath("//m:oMath|//m:oMathPara", namespaces=NS))
    numbered: dict[str, list[int]] = defaultdict(list)
    unnumbered = []
    candidates = []
    for index, paragraph in enumerate(document.paragraphs, 1):
        if index < body_start:
            continue
        text = paragraph_text(paragraph)
        if not looks_like_display_formula(text):
            continue
        candidates.append({"index": index, "text": text[:200], "style": paragraph.style.name})
        match = re.search(r"[（(](\d+)[-－](\d+)[）)]\s*$", text)
        if match:
            numbered[match.group(1)].append(int(match.group(2)))
        else:
            unnumbered.append({"index": index, "text": text[:200]})

    table_candidates = []
    for table_index, table in enumerate(document.tables, 1):
        if len(table.columns) < 3:
            continue
        for row_index, row in enumerate(table.rows, 1):
            number_text = row.cells[-1].text.strip()
            formula_text = " ".join(cell.text.strip() for cell in row.cells[1:-1] if cell.text.strip())
            has_omath = any("m:oMath" in cell._tc.xml or "m:oMathPara" in cell._tc.xml for cell in row.cells)
            if not (has_omath or formula_text or number_text):
                continue
            match = re.search(r"[（(](\d+)[-－](\d+)[）)]\s*$", number_text)
            if not (match or has_omath or looks_like_display_formula(formula_text)):
                continue
            table_candidates.append(
                {
                    "table_index": table_index,
                    "row_index": row_index,
                    "number": number_text,
                    "formula_text": formula_text[:200],
                    "has_omath": has_omath,
                }
            )
            if match:
                numbered[match.group(1)].append(int(match.group(2)))
            elif formula_text:
                unnumbered.append({"index": f"T{table_index}R{row_index}", "text": formula_text[:200]})

    continuity = {}
    for chapter, nums in numbered.items():
        ordered = sorted(nums)
        expected = list(range(1, len(ordered) + 1))
        continuity[chapter] = {"numbers": ordered, "continuous": ordered == expected}

    return {
        "omath_count": omath_count,
        "display_formula_candidates": candidates[:80],
        "display_formula_table_candidates": table_candidates[:80],
        "unnumbered_display_formulas": unnumbered[:80],
        "numbering_by_chapter": continuity,
    }


def build_relationship_map(archive: ZipFile) -> dict[str, str]:
    rel_root = etree.fromstring(archive.read("word/_rels/document.xml.rels"))
    result = {}
    for rel in rel_root.xpath("//rel:Relationship", namespaces=NS):
        result[rel.get("Id")] = rel.get("Target", "")
    return result


def print_section_report(sections: list[dict]) -> None:
    print("## 分节与页码")
    for item in sections:
        print(
            f"- 第 {item['index']} 节：页码格式={item['page_number_format'] or 'None'}，"
            f"起始值={item['page_number_start'] or 'None'}，"
            f"header={item['headers']}，footer={item['footers']}"
        )


def print_toc_report(toc: dict) -> None:
    print("\n## 目录与字段")
    print(f"- TOC 域存在：{toc['has_toc_field']}")
    print(f"- 目录标题段落：{toc['toc_title_index']}")
    print(f"- 目录行数：{toc['line_count']}")
    if toc["repeated_headings"]:
        print(f"- 重复目录条目：{toc['repeated_headings']}")
    else:
        print("- 重复目录条目：无")
    for line in toc["toc_lines"][:25]:
        print(f"  - 第 {line['index']} 段 [{line['style']}] {line['text']}")


def print_header_footer_report(hf: dict) -> None:
    print("\n## 页眉页脚资源")
    for item in hf["headers"]:
        print(
            f"- {item['name']}：drawing={item['has_drawing']} image_rel={item['has_image_rel']} "
            f"table={item['has_table']} underline={item['has_underline']} text={item['text_preview']}"
        )
    for item in hf["footers"]:
        print(
            f"- {item['name']}：PAGE={item['has_page_field']} NUMPAGES={item['has_numpages_field']} text={item['text_preview']}"
        )


def print_formula_report(formulas: dict) -> None:
    print("\n## 公式对象")
    print(f"- OMath 对象数：{formulas['omath_count']}")
    print(f"- 展示公式候选数：{len(formulas['display_formula_candidates'])}")
    print(f"- 表格型公式候选数：{len(formulas['display_formula_table_candidates'])}")
    print(f"- 未编号展示公式数：{len(formulas['unnumbered_display_formulas'])}")
    if formulas["omath_count"] == 0 and (
        formulas["display_formula_candidates"] or formulas["display_formula_table_candidates"]
    ):
        print("- 提示：存在展示公式候选但 OMath 对象数为 0。若已进入定稿阶段，优先用 scripts/repair_formula_objects_with_word.ps1 恢复真正的 Word 公式对象。")
    for chapter, item in sorted(formulas["numbering_by_chapter"].items(), key=lambda x: int(x[0])):
        print(f"- 第 {chapter} 章公式编号：{item['numbers']}，连续={item['continuous']}")
    for item in formulas["display_formula_table_candidates"][:20]:
        print(
            f"  - 表格 T{item['table_index']} 第 {item['row_index']} 行："
            f"number={item['number'] or 'None'} has_omath={item['has_omath']} text={item['formula_text']}"
        )
    for item in formulas["unnumbered_display_formulas"][:20]:
        print(f"  - 第 {item['index']} 段：{item['text']}")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()

    if not args.docx.exists():
        raise SystemExit(f"missing file: {args.docx}")

    document = Document(str(args.docx))
    body_start = detect_body_start(document)
    with ZipFile(args.docx) as archive:
        document_xml_text = read_archive_text(archive, "word/document.xml")
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        rel_map = build_relationship_map(archive)
        sections = analyze_sections(document_xml, rel_map)
        toc = analyze_toc(document, document_xml_text, body_start)
        headers_footers = analyze_headers_and_footers(archive)
        formulas = analyze_formulas(document, document_xml, body_start)

    print(f"# CSU Thesis Word 结构审计：{args.docx.name}")
    print(f"- 段落数：{len(document.paragraphs)}")
    print(f"- 表格数：{len(document.tables)}")
    print(f"- 正文起点（估计）：第 {body_start} 段")
    print_section_report(sections)
    print_toc_report(toc)
    print_header_footer_report(headers_footers)
    print_formula_report(formulas)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
