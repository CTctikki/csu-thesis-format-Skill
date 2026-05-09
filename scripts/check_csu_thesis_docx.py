#!/usr/bin/env python3
"""按中南大学本科毕业论文模板启发式检查 DOCX 排版。"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile

from docx import Document

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def qn(tag: str) -> str:
    return f"{{{W}}}{tag}"


def w_attr(el, name: str = "val"):
    return el.get(qn(name)) if el is not None else None


def approx(actual: float, expected: float, tolerance: float = 0.08) -> bool:
    return abs(actual - expected) <= tolerance


@dataclass
class Finding:
    level: str
    location: str
    message: str


class Report:
    def __init__(self) -> None:
        self.findings: list[Finding] = []

    def add(self, level: str, location: str, message: str) -> None:
        self.findings.append(Finding(level, location, message))

    def error(self, location: str, message: str) -> None:
        self.add("ERROR", location, message)

    def warn(self, location: str, message: str) -> None:
        self.add("WARN", location, message)

    def info(self, location: str, message: str) -> None:
        self.add("INFO", location, message)

    @property
    def error_count(self) -> int:
        return sum(1 for finding in self.findings if finding.level == "ERROR")

    @property
    def warn_count(self) -> int:
        return sum(1 for finding in self.findings if finding.level == "WARN")


def text_of(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def p_pr(paragraph):
    return paragraph._p.pPr


def jc(paragraph) -> str | None:
    props = p_pr(paragraph)
    return w_attr(props.find(qn("jc"))) if props is not None and props.find(qn("jc")) is not None else None


def spacing_multiple(paragraph) -> float | None:
    props = p_pr(paragraph)
    if props is None:
        return None
    spacing = props.find(qn("spacing"))
    if spacing is None:
        return None
    line = spacing.get(qn("line"))
    rule = spacing.get(qn("lineRule"))
    if not line or rule in {"exact", "atLeast"}:
        return None
    return round(int(line) / 240, 3)


def first_line_chars(paragraph) -> int | None:
    props = p_pr(paragraph)
    if props is None:
        return None
    indent = props.find(qn("ind"))
    if indent is None:
        return None
    value = indent.get(qn("firstLineChars"))
    if value is not None:
        return int(value)
    value = indent.get(qn("firstLine"))
    if value is not None:
        return round(int(value) / 240)
    return None


def outline_level(paragraph) -> int | None:
    props = p_pr(paragraph)
    if props is None:
        return None
    outline = props.find(qn("outlineLvl"))
    value = w_attr(outline)
    return int(value) if value is not None else None


def first_run_rpr(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run._r.rPr
    return None


def style_chain(style):
    chain = []
    current = style
    while current is not None:
        chain.append(current)
        current = current.base_style
    return chain


def effective_jc(paragraph) -> str | None:
    props = p_pr(paragraph)
    if props is not None:
        value = w_attr(props.find(qn("jc")))
        if value is not None:
            return value
    for style in style_chain(paragraph.style):
        props = style._element.find(qn("pPr"))
        if props is None:
            continue
        value = w_attr(props.find(qn("jc")))
        if value is not None:
            return value
    return None


def effective_indent(paragraph, attrs: list[str]) -> str | None:
    props = p_pr(paragraph)
    if props is not None:
        indent = props.find(qn("ind"))
        if indent is not None:
            for attr in attrs:
                value = indent.get(qn(attr))
                if value is not None:
                    return value
    for style in style_chain(paragraph.style):
        props = style._element.find(qn("pPr"))
        if props is None:
            continue
        indent = props.find(qn("ind"))
        if indent is None:
            continue
        for attr in attrs:
            value = indent.get(qn(attr))
            if value is not None:
                return value
    return None


def effective_run_rpr(paragraph):
    r_pr = first_run_rpr(paragraph)
    if r_pr is not None:
        return r_pr
    for style in style_chain(paragraph.style):
        r_pr = style._element.find(qn("rPr"))
        if r_pr is not None:
            return r_pr
    return None


def first_run_size(paragraph) -> float | None:
    r_pr = first_run_rpr(paragraph)
    if r_pr is not None:
        size = r_pr.find(qn("sz"))
        value = w_attr(size)
        if value:
            return round(int(value) / 2, 2)
    for style in style_chain(paragraph.style):
        r_pr = style._element.find(qn("rPr"))
        if r_pr is None:
            continue
        size = r_pr.find(qn("sz"))
        value = w_attr(size)
        if value:
            return round(int(value) / 2, 2)
    return None


def first_run_fonts(paragraph) -> set[str]:
    result = set()
    r_pr = first_run_rpr(paragraph)
    if r_pr is not None:
        fonts = r_pr.find(qn("rFonts"))
        if fonts is not None:
            for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
                value = fonts.get(qn(key))
                if value:
                    result.add(value)
    if result:
        return result
    for style in style_chain(paragraph.style):
        r_pr = style._element.find(qn("rPr"))
        if r_pr is None:
            continue
        fonts = r_pr.find(qn("rFonts"))
        if fonts is None:
            continue
        for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
            value = fonts.get(qn(key))
            if value:
                result.add(value)
        if result:
            return result
    return result


def has_font(paragraph, names: list[str]) -> bool:
    fonts = first_run_fonts(paragraph)
    return any(any(name.lower() in font.lower() for font in fonts) for name in names)


def is_bold(paragraph) -> bool | None:
    r_pr = first_run_rpr(paragraph)
    if r_pr is not None:
        bold = r_pr.find(qn("b"))
        if bold is not None:
            return w_attr(bold) not in {"0", "false", "False"}
    for style in style_chain(paragraph.style):
        r_pr = style._element.find(qn("rPr"))
        if r_pr is None:
            continue
        bold = r_pr.find(qn("b"))
        if bold is not None:
            return w_attr(bold) not in {"0", "false", "False"}
    return None


def is_centered(paragraph) -> bool:
    return effective_jc(paragraph) == "center"


def paragraph_location(index: int, text: str) -> str:
    return f"第 {index} 段：{text[:32]}"


def find_body_start(document: Document) -> int:
    """返回正文大概率开始的 1 基段落序号。"""
    candidates = []
    for index, paragraph in enumerate(document.paragraphs, 1):
        text = text_of(paragraph)
        if not re.match(r"^第[0-9一二三四五六七八九十百]+章\s+\S", text):
            continue
        if "\t" in text or re.search(r"\.{2,}\s*\d+$|…+\s*\d+$|\s+\d+$", text):
            continue
        candidates.append(index)
    return candidates[0] if candidates else 1


def find_references_start(document: Document) -> int | None:
    for index, paragraph in enumerate(document.paragraphs, 1):
        if text_of(paragraph) == "参考文献":
            return index
    return None


def docx_member_text(path: Path, member: str) -> str:
    with ZipFile(path) as archive:
        return archive.read(member).decode("utf-8", errors="ignore")


def check_sections(document: Document, report: Report) -> None:
    expected = {
        "page_width": 21.0,
        "page_height": 29.7,
        "top": 2.5,
        "bottom": 2.5,
        "left": 3.0,
        "right": 2.0,
        "footer": 1.75,
    }
    for index, section in enumerate(document.sections, 1):
        location = f"第 {index} 节"
        values = {
            "page_width": section.page_width.cm,
            "page_height": section.page_height.cm,
            "top": section.top_margin.cm,
            "bottom": section.bottom_margin.cm,
            "left": section.left_margin.cm,
            "right": section.right_margin.cm,
            "footer": section.footer_distance.cm,
        }
        labels = {
            "page_width": "页面宽度",
            "page_height": "页面高度",
            "top": "上边距",
            "bottom": "下边距",
            "left": "左边距",
            "right": "右边距",
            "footer": "页脚距边界",
        }
        for key, expected_value in expected.items():
            if not approx(values[key], expected_value):
                report.error(location, f"{labels[key]}为 {values[key]:.2f} cm；模板要求约 {expected_value:.2f} cm")
        header = section.header_distance.cm
        if not (approx(header, 1.25, 0.12) or approx(header, 1.5, 0.12)):
            report.warn(location, f"页眉距边界为 {header:.2f} cm；模板通常约为 1.25 cm 或 1.50 cm")
    if document.sections and not document.sections[0].different_first_page_header_footer:
        report.warn("第 1 节", "封面/前置部分通常应启用“首页不同”的页眉页脚")


def check_toc_and_fields(docx_path: Path, document: Document, report: Report, body_start: int) -> None:
    toc_title = None
    for index, paragraph in enumerate(document.paragraphs, 1):
        if text_of(paragraph) == "目录":
            toc_title = index
            break
    document_xml = docx_member_text(docx_path, "word/document.xml")
    has_toc_field = "TOC " in document_xml
    if toc_title is None:
        report.warn("目录", "未找到“目录”标题")
        return

    toc_lines = []
    for index in range(toc_title + 1, min(body_start, len(document.paragraphs) + 1)):
        text = text_of(document.paragraphs[index - 1])
        if text:
            toc_lines.append(text)

    repeated_ch1 = sum(1 for text in toc_lines if re.match(r"^第1章", text))
    if repeated_ch1 > 1:
        report.warn("目录", "目录中检测到重复的“第1章”起始条目，通常表示静态目录和新目录共存")

    if has_toc_field and toc_lines:
        report.info("目录", "DOCX 中仍包含 TOC 域；如果目标是最终交付版，建议更新后静态定稿")

    title_para = document.paragraphs[toc_title - 1]
    title_size = first_run_size(title_para)
    if not is_centered(title_para):
        report.warn("目录", "“目录”标题应居中")
    if title_size and not approx(title_size, 16, 0.8):
        report.warn("目录", f"“目录”标题字号为 {title_size} pt；常见定稿值约为 16 pt")

    for index in range(toc_title + 1, min(body_start, len(document.paragraphs) + 1)):
        paragraph = document.paragraphs[index - 1]
        text = text_of(paragraph)
        if not text:
            continue
        location = paragraph_location(index, text)
        size = first_run_size(paragraph)
        if size and not (approx(size, 12, 1.0) or approx(size, 10.5, 0.8)):
            report.warn(location, f"目录条目字号为 {size} pt；当前验证过的定稿样式通常为 10.5 pt 或 12 pt")


def check_heading_theme_chain(docx_path: Path, report: Report) -> None:
    styles_xml = docx_member_text(docx_path, "word/styles.xml")
    for style_id in ["1", "21", "10", "22"]:
        pattern = rf'<w:style[^>]*w:styleId="{style_id}".*?</w:style>'
        match = re.search(pattern, styles_xml, re.S)
        if not match:
            continue
        style_xml = match.group(0)
        if 'themeColor="accent1"' in style_xml or 'w:themeColor="accent1"' in style_xml:
            report.warn(f"样式 {style_id}", "标题样式链仍引用 accent1 主题色，预览中可能显示为蓝色而非黑色")


def check_header_assets(docx_path: Path, report: Report) -> None:
    with ZipFile(docx_path) as archive:
        header_rels = [name for name in archive.namelist() if name.startswith("word/_rels/header") and name.endswith(".rels")]
        image_header_rels = 0
        for rel in header_rels:
            rel_xml = archive.read(rel).decode("utf-8", errors="ignore")
            if "relationships/image" in rel_xml:
                image_header_rels += 1
    if image_header_rels == 0:
        report.warn("页眉", "未检测到页眉图片关系；若模板要求左侧校徽校名，当前文件可能只有纯文字页眉")


def looks_like_display_formula(text: str) -> bool:
    if len(text) > 160:
        return False
    if text.startswith(("式中", "其中", "则可通过", "表", "图", "第", "摘要", "目录", "参考文献")):
        return False
    return any(token in text for token in ["=", "Σ", "ρ", "α", "λ", "×", "/"]) and not text.endswith("。")


def check_display_formulas(document: Document, report: Report, body_start: int, references_start: int | None, docx_path: Path) -> None:
    document_xml = docx_member_text(docx_path, "word/document.xml")
    omath_count = document_xml.count("<m:oMath") + document_xml.count("<m:oMathPara")
    formula_numbers: dict[str, list[int]] = {}

    for index, paragraph in enumerate(document.paragraphs, 1):
        if index < body_start:
            continue
        if references_start is not None and index >= references_start:
            continue
        text = text_of(paragraph)
        if not looks_like_display_formula(text):
            continue
        location = paragraph_location(index, text)
        match = re.search(r"[（(](\d+)[-－](\d+)[）)]\s*$", text)
        if match:
            chapter = match.group(1)
            number = int(match.group(2))
            formula_numbers.setdefault(chapter, []).append(number)
        else:
            report.warn(location, "疑似展示型公式未编号，或公式编号不在行末")

    for chapter, values in formula_numbers.items():
        ordered = sorted(values)
        expected = list(range(1, len(ordered) + 1))
        if ordered != expected:
            report.warn(f"第 {chapter} 章公式", f"公式编号为 {ordered}；应连续为 {expected}")

    if omath_count == 0 and formula_numbers:
        report.info("公式", "未检测到 Word OMath 公式对象；当前文档可能使用文本/表格伪公式布局")


def check_title_paragraph(index: int, paragraph, report: Report) -> None:
    text = text_of(paragraph)
    location = paragraph_location(index, text)
    size = first_run_size(paragraph)

    if text in {"摘要", "目录", "参考文献"} or text in {"结束语", "致谢", "结束语(或致谢)"}:
        if not is_centered(paragraph):
            report.warn(location, "前置/后置部分一级标题应居中")
        if size and not approx(size, 16, 0.8):
            report.warn(location, f"标题字号为 {size} pt；模板要求三号，约 16 pt")
        if not has_font(paragraph, ["黑体", "SimHei"]):
            report.warn(location, "标题应使用黑体/SimHei")

    if text == "ABSTRACT":
        if not is_centered(paragraph):
            report.warn(location, "ABSTRACT 标题应居中")
        if size and not approx(size, 16, 0.8):
            report.warn(location, f"ABSTRACT 字号为 {size} pt；模板要求约 16 pt")
        if has_font(paragraph, ["Times New Roman"]) and is_bold(paragraph) is False:
            report.warn(location, "ABSTRACT 应使用加粗 Times New Roman")


def check_body_heading(index: int, paragraph, report: Report, body_start: int) -> None:
    if index < body_start:
        return
    text = text_of(paragraph)
    location = paragraph_location(index, text)
    size = first_run_size(paragraph)

    if re.match(r"^第[0-9一二三四五六七八九十百]+章\s+\S", text):
        if not is_centered(paragraph):
            report.warn(location, "章标题应居中")
        if size and not approx(size, 16, 0.8):
            report.warn(location, f"章标题字号为 {size} pt；模板要求约 16 pt")
        if not has_font(paragraph, ["黑体", "SimHei"]):
            report.warn(location, "章标题应使用黑体/SimHei")
        return

    if re.match(r"^\d+\.\d+\.\d+\s+\S", text):
        if size and not approx(size, 12, 0.8):
            report.warn(location, f"三级标题字号为 {size} pt；模板要求 12 pt")
        if not has_font(paragraph, ["楷", "Kai"]):
            report.warn(location, "三级标题应使用楷体/楷体GB2312")
        chars = first_line_chars(paragraph)
        if chars not in {None, 200, 2}:
            report.warn(location, "三级标题应缩进约两个汉字")
        return

    if re.match(r"^\d+\.\d+\s+\S", text):
        if size and not approx(size, 12, 0.8):
            report.warn(location, f"二级标题字号为 {size} pt；模板要求 12 pt")
        if not has_font(paragraph, ["黑体", "SimHei"]):
            report.warn(location, "二级标题应使用黑体/SimHei")
        chars = first_line_chars(paragraph)
        if chars not in {None, 200, 2}:
            report.warn(location, "二级标题应缩进约两个汉字")


def check_caption_or_equation(index: int, paragraph, report: Report, body_start: int) -> None:
    if index < body_start:
        return
    text = text_of(paragraph)
    location = paragraph_location(index, text)
    size = first_run_size(paragraph)

    if re.match(r"^表\d+[-－]\d+\s+\S", text):
        if not is_centered(paragraph):
            report.warn(location, "表题应位于表格上方并居中")
        if size and not approx(size, 10.5, 0.8):
            report.warn(location, f"表题字号为 {size} pt；模板要求五号，约 10.5 pt")
        if not has_font(paragraph, ["黑体", "SimHei"]):
            report.warn(location, "表题应使用黑体/SimHei")

    if re.match(r"^图\d+[-－]\d+\s+\S", text):
        if not is_centered(paragraph):
            report.warn(location, "图题应位于图片下方并居中")
        if size and not approx(size, 10.5, 0.8):
            report.warn(location, f"图题字号为 {size} pt；模板要求五号，约 10.5 pt")
        if not has_font(paragraph, ["黑体", "SimHei"]):
            report.warn(location, "图题应使用黑体/SimHei")

    if re.fullmatch(r"[（(]\d+[-－]\d+[）)]", text):
        if jc(paragraph) not in {"right", "end"}:
            report.warn(location, "单独一行的公式编号应右对齐")


def looks_like_body(text: str) -> bool:
    if len(text) < 25:
        return False
    if re.match(r"^(第[0-9一二三四五六七八九十百]+章|\d+\.\d+|图\d+|表\d+|摘要|ABSTRACT|目录|参考文献|结束语|致谢)", text):
        return False
    if re.fullmatch(r"[（(]\d+[-－]\d+[）)]", text):
        return False
    return True


def check_body_paragraphs(document: Document, report: Report, body_start: int, references_start: int | None, limit: int = 25) -> None:
    seen = 0
    for index, paragraph in enumerate(document.paragraphs, 1):
        if index < body_start:
            continue
        if references_start is not None and index >= references_start:
            continue
        text = text_of(paragraph)
        if not looks_like_body(text):
            continue
        seen += 1
        if seen > limit:
            return
        location = paragraph_location(index, text)
        line = spacing_multiple(paragraph)
        if line is not None and not approx(line, 1.5, 0.05):
            report.warn(location, f"正文段落行距为 {line}；模板要求 1.5 倍行距")
        chars = first_line_chars(paragraph)
        if chars not in {None, 200, 2}:
            report.warn(location, "正文段落应首行缩进约两个汉字")
        size = first_run_size(paragraph)
        if size and not approx(size, 12, 1.0):
            report.warn(location, f"正文字号为 {size} pt；模板要求小四号，约 12 pt")


def check_keywords(document: Document, report: Report) -> None:
    for index, paragraph in enumerate(document.paragraphs, 1):
        text = text_of(paragraph)
        if text.startswith("关键词"):
            terms = re.split(r"\s{2,}|　|；|;|，|,", text.split("：", 1)[-1].strip())
            terms = [term for term in terms if term]
            if not (3 <= len(terms) <= 8):
                report.warn(paragraph_location(index, text), f"中文关键词数量为 {len(terms)}；模板要求 3-8 个")
        if text.lower().startswith(("key words", "keywords")):
            terms = re.split(r"\s{2,}|;|,|；|，", re.split(r"[:：]", text, 1)[-1].strip())
            terms = [term for term in terms if term]
            if terms and not (3 <= len(terms) <= 8):
                report.warn(paragraph_location(index, text), f"英文关键词数量为 {len(terms)}；模板要求 3-8 个")


def check_references(document: Document, report: Report) -> None:
    in_refs = False
    numbers: list[int] = []
    for index, paragraph in enumerate(document.paragraphs, 1):
        text = text_of(paragraph)
        if text == "参考文献":
            in_refs = True
            continue
        if not in_refs:
            continue
        match = re.match(r"^\[(\d+)\]\s+", text)
        if match:
            numbers.append(int(match.group(1)))
            size = first_run_size(paragraph)
            if size and not approx(size, 10.5, 1.0):
                report.warn(paragraph_location(index, text), f"参考文献条目字号为 {size} pt；模板要求五号，约 10.5 pt")
            if not has_font(paragraph, ["楷", "Kai"]):
                report.warn(paragraph_location(index, text), "参考文献条目应使用楷体/楷体GB2312")
            left = effective_indent(paragraph, ["left"])
            hanging = effective_indent(paragraph, ["hanging"])
            if left != "420" or hanging != "420":
                report.warn(paragraph_location(index, text), "参考文献条目通常应使用悬挂缩进 420/420")
    if numbers:
        expected = list(range(1, len(numbers) + 1))
        if numbers != expected:
            report.warn("参考文献", f"参考文献编号为 {numbers}；应按出现顺序连续编号为 {expected}")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--strict", action="store_true", help="发现 ERROR 时以非零状态码退出")
    args = parser.parse_args()

    if not args.docx.exists():
        raise SystemExit(f"找不到文件：{args.docx}")

    document = Document(str(args.docx))
    report = Report()
    body_start = find_body_start(document)
    references_start = find_references_start(document)

    check_sections(document, report)
    check_toc_and_fields(args.docx, document, report, body_start)
    check_heading_theme_chain(args.docx, report)
    check_header_assets(args.docx, report)
    check_display_formulas(document, report, body_start, references_start, args.docx)
    for index, paragraph in enumerate(document.paragraphs, 1):
        if not text_of(paragraph):
            continue
        check_title_paragraph(index, paragraph, report)
        check_body_heading(index, paragraph, report, body_start)
        check_caption_or_equation(index, paragraph, report, body_start)
    check_body_paragraphs(document, report, body_start, references_start)
    check_keywords(document, report)
    check_references(document, report)

    print(f"# 中南大学论文 DOCX 检查报告：{args.docx.name}\n")
    print("## 汇总")
    print(f"- 错误：{report.error_count}")
    print(f"- 提醒：{report.warn_count}")
    print(f"- 已检查段落数：{len(document.paragraphs)}")
    print(f"- 已检查节数：{len(document.sections)}")
    print(f"- 发现表格数：{len(document.tables)}")
    print(f"- 正文大约从第 {body_start} 段开始")

    print("\n## 问题清单")
    if not report.findings:
        print("- 启发式检查未发现明显的中南大学模板格式问题。")
    else:
        level_labels = {"ERROR": "错误", "WARN": "提醒", "INFO": "信息"}
        for finding in report.findings:
            print(f"- [{level_labels.get(finding.level, finding.level)}] {finding.location}：{finding.message}")

    print("\n## 提醒")
    print("- 本脚本为启发式检查。最终请将 DOCX 渲染为 PDF，重点目视检查封面、摘要、目录、每章首页、题注和参考文献。")

    return 1 if args.strict and report.error_count else 0


if __name__ == "__main__":
    raise SystemExit(main())
